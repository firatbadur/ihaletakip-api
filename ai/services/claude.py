"""
Claude (Anthropic) analiz servisi.

Firebase Cloud Function'daki analiz mantığının Django karşılığı.
Doküman (PDF/DOCX/DOC/TXT) metnini çıkarır, prompt'u seçer ve Claude'a gönderir.

İki latent bug düzeltildi:
  1) APIConnectionError/APITimeoutError'da status_code yok → güvenli getattr.
  2) Eski ikili .doc formatı → net "desteklenmiyor" hatası (sessizce
     bozuk metin göndermek yerine).
"""
import base64
import io
import json
import logging

from django.conf import settings

logger = logging.getLogger("ihaletakip")

SUPPORTED_EXTENSIONS = ["pdf", "docx", "doc", "txt"]


class AnalysisError(Exception):
    """İş kuralı hatası (istemciye mesaj döner). status varsayılan 400."""

    def __init__(self, message, status=400):
        super().__init__(message)
        self.message = message
        self.status = status


# ── Anahtar / model ────────────────────────────────────
def get_api_key() -> str:
    """
    Claude API anahtarını al: önce env (ANTHROPIC_API_KEY),
    yoksa AppSetting('anthropic_api_key') — Firestore config/ai_service karşılığı.
    """
    key = settings.ANTHROPIC_API_KEY
    if not key:
        from core.models import AppSetting

        key = AppSetting.get("anthropic_api_key")
    if not key:
        raise AnalysisError("AI servisi yapılandırılmamış (API anahtarı yok).", status=503)
    return key


# ── Doküman metin çıkarımı ─────────────────────────────
def extract_text_from_docx(file_bytes: bytes) -> str:
    from docx import Document as DocxDocument

    doc = DocxDocument(io.BytesIO(file_bytes))
    parts = [p.text.strip() for p in doc.paragraphs if p.text.strip()]
    for table in doc.tables:
        for row in table.rows:
            row_text = " | ".join(c.text.strip() for c in row.cells if c.text.strip())
            if row_text:
                parts.append(row_text)
    return "\n".join(parts)


def extract_text_from_pdf(file_bytes: bytes) -> str:
    from pypdf import PdfReader

    reader = PdfReader(io.BytesIO(file_bytes))
    pages = []
    for page in reader.pages:
        text = page.extract_text()
        if text:
            pages.append(text.strip())
    return "\n\n".join(pages)


def build_document_block(file_base64: str, file_name: str) -> dict:
    """
    Dosyayı çözüp uygun içerik bloğunu üretir.
    PDF → doğrudan Claude'a document bloğu; diğerleri → metin bloğu.
    """
    if not file_base64 or not file_name:
        raise AnalysisError("Dosya veya dosya adı eksik.")

    extension = file_name.rsplit(".", 1)[-1].lower() if "." in file_name else ""

    try:
        file_bytes = base64.b64decode(file_base64)
    except Exception as e:
        raise AnalysisError("Dosya çözümlenemedi.") from e

    if len(file_bytes) > settings.MAX_FILE_SIZE:
        mb = len(file_bytes) / (1024 * 1024)
        raise AnalysisError(f"Dosya çok büyük ({mb:.1f} MB). Maksimum 10 MB.")

    if extension == "pdf":
        return {
            "type": "document",
            "source": {
                "type": "base64",
                "media_type": "application/pdf",
                "data": file_base64,
            },
        }

    if extension == "docx":
        text = extract_text_from_docx(file_bytes)
        if not text.strip():
            raise AnalysisError("Doküman içeriği boş veya okunamadı.")
        return {"type": "text", "text": f"[Doküman: {file_name}]\n\n{text[:100000]}"}

    if extension == "txt":
        text = file_bytes.decode("utf-8", errors="ignore")
        return {"type": "text", "text": f"[Doküman: {file_name}]\n\n{text[:100000]}"}

    if extension == "doc":
        # Eski ikili .doc güvenilir parse edilemez → net hata döndür.
        raise AnalysisError(
            "Eski Word (.doc) formatı desteklenmiyor. "
            "Lütfen dosyayı PDF veya .docx olarak yükleyin."
        )

    raise AnalysisError(f".{extension} desteklenmiyor. PDF, DOCX veya TXT kullanın.")


# ── Claude çağrısı ─────────────────────────────────────
def call_claude(api_key: str, content_blocks: list, prompt: str, max_tokens: int = None) -> dict:
    """Claude API'ye istek gönder; {analysis, usage} döndür."""
    import anthropic

    client = anthropic.Anthropic(api_key=api_key)
    content_blocks = list(content_blocks) + [{"type": "text", "text": prompt}]

    try:
        message = client.messages.create(
            model=settings.CLAUDE_MODEL,
            max_tokens=max_tokens or settings.CLAUDE_MAX_TOKENS,
            messages=[{"role": "user", "content": content_blocks}],
        )
    except anthropic.APIError as e:
        # BUG FIX: bağlantı/timeout hatalarında status_code yoktur.
        code = getattr(e, "status_code", None) or "bağlantı"
        logger.warning("Claude API hatası (%s): %s", code, e)
        raise AnalysisError(f"AI servisi geçici olarak yanıt vermiyor ({code}).", status=502) from e

    analysis = "".join(b.text for b in message.content if getattr(b, "type", "") == "text")
    return {
        "analysis": analysis,
        "usage": {
            "input_tokens": message.usage.input_tokens,
            "output_tokens": message.usage.output_tokens,
        },
    }


# ── Prompt oluşturucular ───────────────────────────────
def _keywords_meta_text(tender_meta: dict) -> str:
    t = tender_meta or {}
    lines = ["\n\nİhale Bilgileri:"]
    lines.append(f"- İhale Adı: {t.get('ihaleAdi', 'Bilinmiyor')}")
    lines.append(f"- İdare: {t.get('idareAdi', 'Bilinmiyor')}")
    lines.append(f"- İl: {t.get('il', 'Bilinmiyor')}")
    if t.get("okasNames"):
        lines.append(f"- OKAS Kategorileri: {', '.join(t['okasNames'])}")
    if t.get("ihaleTip"):
        tip_map = {"1": "Mal Alımı", "2": "Yapım", "3": "Hizmet", "4": "Danışmanlık"}
        lines.append(f"- İhale Türü: {tip_map.get(str(t['ihaleTip']), str(t['ihaleTip']))}")
    return "\n".join(lines)


def _cost_analysis_prompt(base_prompt: str, tender_meta: dict, similar_tenders: list) -> str:
    prompt = base_prompt
    if tender_meta:
        t = tender_meta
        meta = ["\n\n## MEVCUT İHALE BİLGİLERİ:"]
        meta.append(f"- İhale Adı: {t.get('ihaleAdi', 'Bilinmiyor')}")
        meta.append(f"- İdare: {t.get('idareAdi', 'Bilinmiyor')}")
        meta.append(f"- İl: {t.get('il', 'Bilinmiyor')}")
        if t.get("okasNames"):
            meta.append(f"- OKAS: {', '.join(t['okasNames'])}")
        prompt += "\n".join(meta)
    if similar_tenders:
        block = "\n\n## BENZER İHALE VERİLERİ (Referans):\n"
        block += "Aşağıda aynı türde sonuçlanmış benzer ihalelerin bilgileri yer almaktadır. "
        block += "Bu verileri piyasa karşılaştırması ve teklif stratejisi için kullan:\n\n"
        block += json.dumps(similar_tenders, ensure_ascii=False, indent=2)
        prompt += block
    return prompt


# ── Ana giriş noktası ──────────────────────────────────
def run_analysis(analysis_type, file_base64=None, file_name=None,
                 tender_meta=None, similar_tenders=None) -> dict:
    """
    Analiz tipine göre Claude çağrısını yürütür.
    Dönen: {analysis, usage}
    Hata: AnalysisError fırlatır.
    """
    from ai.prompts import PROMPTS

    prompt = PROMPTS.get(analysis_type)
    if not prompt:
        raise AnalysisError("Geçersiz analiz türü.")

    api_key = get_api_key()

    if analysis_type == "generate_keywords":
        if not tender_meta:
            raise AnalysisError("İhale bilgileri eksik.")
        full_prompt = prompt + _keywords_meta_text(tender_meta)
        return call_claude(api_key, [], full_prompt, max_tokens=150)

    if analysis_type == "cost_analysis":
        full_prompt = _cost_analysis_prompt(prompt, tender_meta, similar_tenders)
        return call_claude(api_key, [], full_prompt)

    # tech_spec / admin_spec → dosya gerekli
    block = build_document_block(file_base64, file_name)
    return call_claude(api_key, [block], prompt)
